import sys
import re
import os
import base64
import hashlib
from datetime import datetime, timedelta
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
                             QLabel, QLineEdit, QComboBox, QPushButton, QFileDialog, QMessageBox, QFrame, QGridLayout,
                             QScrollArea, QListWidget, QListWidgetItem, QInputDialog)
from PyQt6.QtCore import Qt, QEvent, QSize, QPropertyAnimation, QEasingCurve
from docx import Document

# --- 0. 安全与配置逻辑 ---
CONFIG_FILE = "app_settings.bin"


def get_hash(text):
    return hashlib.sha256(text.encode()).hexdigest()


def save_config(status, password_hash):
    content = f"{status}|{password_hash}"
    encoded_content = base64.b64encode(content.encode()).decode()
    with open(CONFIG_FILE, "w") as f:
        f.write(encoded_content)


def load_config():
    if not os.path.exists(CONFIG_FILE):
        return "locked", get_hash("123")
    try:
        with open(CONFIG_FILE, "r") as f:
            encoded_content = f.read()
            decoded = base64.b64decode(encoded_content).decode().split("|")
            return decoded[0], decoded[1]
    except:
        return "locked", get_hash("123")


# --- 1. 高端保密局登录界面 ---
class LoginWindow(QWidget):
    def __init__(self, stored_hash):
        super().__init__()
        self.stored_hash = stored_hash
        self.authorized = False
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint)  # 无边框
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)  # 透明背景支持圆角
        self.setFixedSize(400, 300)
        self.init_ui()

    def init_ui(self):
        # 外层发光边框容器
        self.container = QFrame(self)
        self.container.setObjectName("LoginContainer")
        self.container.setFixedSize(400, 300)

        layout = QVBoxLayout(self.container)
        layout.setContentsMargins(30, 30, 30, 30)
        layout.setSpacing(20)

        # 顶部标题
        title = QLabel("ACCESS RESTRICTED")
        title.setObjectName("LoginTitle")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)

        # 装饰横线
        line = QFrame()
        line.setObjectName("DecorLine")
        line.setFixedHeight(2)

        # 密码输入框
        self.pwd_input = QLineEdit()
        self.pwd_input.setPlaceholderText("ENTER CLEARANCE CODE")
        self.pwd_input.setEchoMode(QLineEdit.EchoMode.Password)
        self.pwd_input.setObjectName("LoginInput")
        self.pwd_input.returnPressed.connect(self.check_auth)

        # 按钮
        btn_layout = QHBoxLayout()
        self.btn_auth = QPushButton("VERIFY")
        self.btn_auth.setObjectName("VerifyBtn")
        self.btn_auth.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_auth.clicked.connect(self.check_auth)

        self.btn_close = QPushButton("ABORT")
        self.btn_close.setObjectName("AbortBtn")
        self.btn_close.clicked.connect(sys.exit)

        btn_layout.addWidget(self.btn_close)
        btn_layout.addWidget(self.btn_auth)

        layout.addWidget(title)
        layout.addWidget(line)
        layout.addStretch()
        layout.addWidget(self.pwd_input)
        layout.addStretch()
        layout.addLayout(btn_layout)

        # 样式表
        self.setStyleSheet("""
            #LoginContainer {
                background-color: #0D1117;
                border: 2px solid #00F0FF;
                border-radius: 15px;
            }
            #LoginTitle {
                color: #00F0FF;
                font-family: 'Courier New';
                font-size: 18px;
                font-weight: bold;
                letter-spacing: 3px;
            }
            #DecorLine {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 rgba(0,240,255,0), stop:0.5 rgba(0,240,255,255), stop:1 rgba(0,240,255,0));
            }
            #LoginInput {
                background-color: #161B22;
                border: 1px solid #30363D;
                border-radius: 5px;
                padding: 10px;
                color: #00F0FF;
                font-family: 'Consolas';
                font-size: 14px;
            }
            #LoginInput:focus {
                border-color: #00F0FF;
            }
            QPushButton {
                font-family: 'Courier New';
                font-weight: bold;
                padding: 8px;
                border-radius: 4px;
            }
            #VerifyBtn {
                background-color: #00F0FF;
                color: #000000;
            }
            #VerifyBtn:hover {
                background-color: #B2FAFF;
            }
            #AbortBtn {
                background-color: transparent;
                border: 1px solid #F85149;
                color: #F85149;
            }
            #AbortBtn:hover {
                background-color: #F85149;
                color: white;
            }
        """)

    def check_auth(self):
        pwd = self.pwd_input.text()
        if get_hash(pwd) == self.stored_hash:
            save_config("unlocked", self.stored_hash)
            self.authorized = True
            self.close()
        else:
            self.pwd_input.clear()
            self.pwd_input.setPlaceholderText("INVALID CODE - ACCESS DENIED")
            self.pwd_input.setStyleSheet("border-color: #F85149; color: #F85149;")


# --- 2. 核心下拉框组件 (拦截滚轮) ---
class MultiSelectComboBox(QComboBox):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setEditable(True)
        self.lineEdit().setReadOnly(True)
        self.lineEdit().setPlaceholderText("可多选...")
        self.lineEdit().setAttribute(Qt.WidgetAttribute.WA_TransparentForMouseEvents)
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.list_widget = QListWidget()
        self.setModel(self.list_widget.model())
        self.setView(self.list_widget)
        self.view().viewport().installEventFilter(self)
        self.model().dataChanged.connect(self.update_text)

    def eventFilter(self, widget, event):
        if widget == self.view().viewport() and event.type() == QEvent.Type.MouseButtonRelease:
            index = self.view().indexAt(event.pos())
            item = self.list_widget.item(index.row())
            if item:
                item.setCheckState(
                    Qt.CheckState.Unchecked if item.checkState() == Qt.CheckState.Checked else Qt.CheckState.Checked)
            return True
        return super().eventFilter(widget, event)

    def mouseReleaseEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            if self.view().isVisible():
                self.hidePopup()
            else:
                self.showPopup()
        else:
            super().mouseReleaseEvent(event)

    def wheelEvent(self, e):
        e.ignore()

    def update_text(self):
        selected = [self.list_widget.item(i).text() for i in range(self.list_widget.count()) if
                    self.list_widget.item(i).checkState() == Qt.CheckState.Checked]
        self.lineEdit().setText(", ".join(selected))

    def currentText(self):
        return self.lineEdit().text()


class NoWheelComboBox(QComboBox):
    def wheelEvent(self, e):
        if not self.view().isVisible():
            e.ignore()
        else:
            super().wheelEvent(e)


# --- 3. 主界面 ---
class MedicalApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("口腔种植病历生成器 v2.8")
        self.init_ui()
        self.apply_dark_style()
        self.showFullScreen()

    def init_ui(self):
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.Shape.NoFrame)
        self.setCentralWidget(scroll)

        main_bg = QWidget()
        scroll.setWidget(main_bg)
        global_layout = QVBoxLayout(main_bg)
        global_layout.setAlignment(Qt.AlignmentFlag.AlignHCenter)

        content_wrapper = QFrame()
        content_wrapper.setFixedWidth(850)
        global_layout.addWidget(content_wrapper)

        layout = QVBoxLayout(content_wrapper)
        layout.setSpacing(20)
        layout.setContentsMargins(25, 25, 25, 25)

        # --- 胶囊式 Header ---
        header_main_layout = QHBoxLayout()
        header_text_vbox = QVBoxLayout()
        title = QLabel("口腔种植病历模板生成系统")
        title.setObjectName("MainTitle")
        subtitle = QLabel("3D Automation Engine | 自动化病历填充引擎")
        subtitle.setObjectName("SubTitle")
        header_text_vbox.addWidget(title)
        header_text_vbox.addWidget(subtitle)
        header_main_layout.addLayout(header_text_vbox)
        header_main_layout.addStretch()

        capsule_frame = QFrame();
        capsule_frame.setObjectName("CapsuleFrame")
        capsule_layout = QHBoxLayout(capsule_frame)
        capsule_layout.setContentsMargins(8, 2, 8, 2);
        capsule_layout.setSpacing(0)

        self.btn_modify_pwd = QPushButton("修改密码");
        self.btn_modify_pwd.setObjectName("GhostBtn_Purple")
        self.btn_modify_pwd.setFixedSize(70, 26);
        self.btn_modify_pwd.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_modify_pwd.clicked.connect(self.change_password)

        line = QFrame();
        line.setFixedWidth(1);
        line.setFixedHeight(14);
        line.setStyleSheet("background-color: #3A3A4A;")

        self.btn_manual_lock = QPushButton("上锁");
        self.btn_manual_lock.setObjectName("GhostBtn_Green")
        self.btn_manual_lock.setFixedSize(50, 26);
        self.btn_manual_lock.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_manual_lock.clicked.connect(self.manual_lock)

        capsule_layout.addWidget(self.btn_modify_pwd);
        capsule_layout.addWidget(line);
        capsule_layout.addWidget(self.btn_manual_lock)
        header_main_layout.addWidget(capsule_frame, 0, Qt.AlignmentFlag.AlignTop)
        layout.addLayout(header_main_layout)

        # --- 表单内容 (逻辑保持不变) ---
        file_section = QFrame();
        file_section.setObjectName("SectionCard")
        file_grid = QGridLayout(file_section)
        self.template_path = QLineEdit();
        self.template_path.setReadOnly(True)
        btn_browse = QPushButton("选择模板");
        btn_browse.setFixedSize(95, 35);
        btn_browse.clicked.connect(self.get_template)
        self.output_name = QLineEdit("生成_口腔种植病历.docx")
        file_grid.addWidget(QLabel("DOCX 模板:"), 0, 0);
        file_grid.addWidget(self.template_path, 0, 1);
        file_grid.addWidget(btn_browse, 0, 2)
        file_grid.addWidget(QLabel("保存名称:"), 1, 0);
        file_grid.addWidget(self.output_name, 1, 1, 1, 2)
        layout.addWidget(file_section)

        form_section = QFrame();
        form_section.setObjectName("SectionCard")
        form_grid = QGridLayout(form_section);
        form_grid.setContentsMargins(25, 30, 25, 30);
        form_grid.setSpacing(20)
        self.inputs = {}
        self.add_form_item(form_grid, "1. 病情描述", "desc", "例如：牙缺失一年以上", 0, 0, 1, 2)
        brands = ["卡瓦盛邦", "士卓曼(亲水)", "士卓曼(非亲水)", "韩国DIO", "Nobel", "Nobel cc", "Nobel pmc",
                  "Nobel pcc", "Nobel active"]
        self.add_multi_combo_item(form_grid, "2. 植体品牌 (多选)", "brand", brands, 1, 0)
        self.date_input = QLineEdit();
        self.date_input.setPlaceholderText("YYYY-MM-DD");
        self.date_input.textChanged.connect(self.format_date)
        self.add_custom_item(form_grid, "4. 种植日期", self.date_input, "date", 1, 1)
        self.add_form_item(form_grid, "3-1. 牙位", "tooth_pos", "如: 46", 2, 0)
        self.add_tooth_model_item(form_grid, "3-2. 植体型号 ", "tooth_model", "输入数字后 按回车", 2, 1)
        self.add_multi_combo_item(form_grid, "5. 颌位选择 (多选)", "jaw", ["上", "下"], 3, 0)
        self.add_multi_combo_item(form_grid, "6. 手术类型 (多选)", "op_type", ["种植", "上颌窦内提升", "上颌窦外提升"],
                                  3, 1)
        self.add_combo_item(form_grid, "7. 种植体个数", "count", [str(i) for i in range(1, 21)], 4, 0)
        self.add_multi_combo_item(form_grid, "8. 冠桥类型 (多选)", "bridge",
                                  ["单冠", "连冠", "上半口桥架", "下半口桥架", "全口桥架"], 4, 1)
        self.add_combo_item(form_grid, "9_1. 高压 (mmHg)", "h_pressure", [str(i) for i in range(110, 141)], 5, 0)
        self.add_combo_item(form_grid, "9_2. 低压 (mmHg)", "l_pressure", [str(i) for i in range(70, 91)], 5, 1)
        self.add_combo_item(form_grid, "10. 心率 (次/分)", "rate", [str(i) for i in range(60, 101)], 6, 0)
        self.s_time_input = QLineEdit();
        self.s_time_input.setPlaceholderText("HH:mm");
        self.s_time_input.textChanged.connect(lambda t: self.format_time(self.s_time_input, t))
        self.add_custom_item(form_grid, "11_1. 开始时间", self.s_time_input, "s_time", 6, 1)
        self.e_time_input = QLineEdit();
        self.e_time_input.setPlaceholderText("HH:mm");
        self.e_time_input.textChanged.connect(lambda t: self.format_time(self.e_time_input, t))
        self.add_custom_item(form_grid, "11_2. 结束时间", self.e_time_input, "e_time", 7, 0)
        self.add_combo_item(form_grid, "12. 牙龈厚度", "thickness", [f"{i}mm" for i in range(1, 11)], 7, 1)
        self.add_combo_item(form_grid, "13. 骨质分类", "bone", ["I 类骨", "II 类骨", "III 类骨", "IV 类骨"], 8, 0, 1, 2)
        layout.addWidget(form_section)

        btn_container = QWidget();
        btn_layout = QHBoxLayout(btn_container)
        self.btn_exit = QPushButton("退出程序");
        self.btn_exit.setObjectName("ExitButton");
        self.btn_exit.setFixedSize(140, 50);
        self.btn_exit.clicked.connect(self.close)
        btn_generate = QPushButton("生成 Word 文档");
        btn_generate.setFixedHeight(50);
        btn_generate.setObjectName("GenerateButton");
        btn_generate.clicked.connect(self.generate_word)
        btn_layout.addWidget(self.btn_exit);
        btn_layout.addSpacing(20);
        btn_layout.addWidget(btn_generate)
        layout.addWidget(btn_container)

    # --- 逻辑 ---
    def keyPressEvent(self, event):
        if event.key() == Qt.Key.Key_Escape:
            if self.isFullScreen():
                self.showNormal();
                self.resize(950, 850)
                screen = self.screen().availableGeometry();
                self.move(int((screen.width() - 950) / 2), int((screen.height() - 850) / 2))
            else:
                self.showFullScreen()
        else:
            super().keyPressEvent(event)

    def manual_lock(self):
        _, h = load_config();
        save_config("locked", h);
        QMessageBox.information(self, "提示", "已上锁")

    def change_password(self):
        p, ok = QInputDialog.getText(self, "安全", "新密码:", QLineEdit.EchoMode.Password)
        if ok and p:
            c, ok2 = QInputDialog.getText(self, "安全", "确认新密码:", QLineEdit.EchoMode.Password)
            if ok2 and p == c: save_config("unlocked", get_hash(p)); QMessageBox.information(self, "成功", "已更新")

    def add_tooth_model_item(self, g, l, k, p, r, c):
        v = QVBoxLayout();
        v.setSpacing(8);
        e = QLineEdit();
        e.setPlaceholderText(p)
        e.returnPressed.connect(lambda: self.handle_model_logic(e));
        v.addWidget(QLabel(l));
        v.addWidget(e)
        self.inputs[k] = e;
        g.addLayout(v, r, c)

    def handle_model_logic(self, edit):
        raw = edit.text().strip();
        match = re.search(r'[\d.]+$', raw)
        if not match: return
        num = match.group(0);
        pre = raw[:raw.rfind(num)].strip()
        if not pre:
            nt = f"{num}mm×"
        elif pre.endswith("×"):
            nt = f"{raw}mm"
        elif pre.endswith("mm"):
            nt = f"{pre.rstrip(',')}, {num}mm×"
        else:
            nt = f"{raw}mm×"
        edit.setText(nt);
        edit.setCursorPosition(len(nt))

    def add_form_item(self, g, l, k, p, r, c, rs=1, cs=1):
        v = QVBoxLayout();
        e = QLineEdit();
        e.setPlaceholderText(p);
        v.setSpacing(8);
        v.addWidget(QLabel(l));
        v.addWidget(e)
        self.inputs[k] = e;
        g.addLayout(v, r, c, rs, cs)

    def add_combo_item(self, g, l, k, items, r, c, rs=1, cs=1):
        v = QVBoxLayout();
        cb = NoWheelComboBox();
        cb.addItems(items);
        v.setSpacing(8);
        v.addWidget(QLabel(l));
        v.addWidget(cb)
        self.inputs[k] = cb;
        g.addLayout(v, r, c, rs, cs)

    def add_multi_combo_item(self, g, l, k, items, r, c, rs=1, cs=1):
        v = QVBoxLayout();
        cb = MultiSelectComboBox();
        cb.addItems(items);
        v.setSpacing(8);
        v.addWidget(QLabel(l));
        v.addWidget(cb)
        self.inputs[k] = cb;
        g.addLayout(v, r, c, rs, cs)

    def add_custom_item(self, g, l, w, k, r, c):
        v = QVBoxLayout();
        v.setSpacing(8);
        v.addWidget(QLabel(l));
        v.addWidget(w);
        self.inputs[k] = w;
        g.addLayout(v, r, c)

    def format_date(self, t):
        self.date_input.blockSignals(True);
        raw = t.replace("-", "")[:8]
        res = f"{raw[:4]}-{raw[4:6]}-{raw[6:8]}".strip("-") if len(raw) > 4 else raw
        self.date_input.setText(res);
        self.date_input.blockSignals(False)

    def format_time(self, w, t):
        w.blockSignals(True);
        raw = t.replace(":", "")[:4]
        res = f"{raw[:2]}:{raw[2:4]}".strip(":") if len(raw) > 2 else raw
        w.setText(res);
        w.blockSignals(False)

    def get_template(self):
        f, _ = QFileDialog.getOpenFileName(self, "模板", "", "Word Files (*.docx)")
        if f: self.template_path.setText(f)

    def generate_word(self):
        p = self.template_path.text()
        if not p: QMessageBox.warning(self, "错误", "请先选择模板！"); return
        try:
            doc = Document(p)
            vals = {k: (v.currentText() if isinstance(v, QComboBox) else v.text()) for k, v in self.inputs.items()}
            rmap = {"{{description}}": vals['desc'], "{{brand}}": vals['brand'], "{{y_position}}": vals['tooth_pos'],
                    "{{model}}": vals['tooth_model'], "{{date}}": vals['date'], "{{e_position}}": vals['jaw'],
                    "{{surgery}}": vals['op_type'], "{{number}}": vals['count'], "{{Crown}}": vals['bridge'],
                    "{{h_pressure}}": vals['h_pressure'], "{{l_pressure}}": vals['l_pressure'],
                    "{{rate}}": vals['rate'],
                    "{{s_time}}": vals['s_time'], "{{o_time}}": vals['e_time'], "{{thickness}}": vals['thickness'],
                    "{{bone}}": vals['bone']}
            for pa in doc.paragraphs:
                for k, v in rmap.items():
                    if k in pa.text: pa.text = pa.text.replace(k, v)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for k, v in rmap.items():
                            if k in cell.text: cell.text = cell.text.replace(k, v)
            sp, _ = QFileDialog.getSaveFileName(self, "保存", self.output_name.text(), "Word Files (*.docx)")
            if sp: doc.save(sp); QMessageBox.information(self, "成功", "已保存")
        except Exception as e:
            QMessageBox.critical(self, "错误", str(e))

    def apply_dark_style(self):
        self.setStyleSheet("""
            QMainWindow, QWidget { background-color: #1A1A21; font-family: "Segoe UI", "Microsoft YaHei UI"; }
            #MainTitle { color: #FFFFFF; font-size: 26px; font-weight: 800; }
            #SubTitle { color: #6A6A80; font-size: 13px; }
            #SectionCard { background-color: #242430; border-radius: 12px; border: 1px solid #3A3A4A; }
            QLabel { color: #A0A0B5; font-size: 13px; font-weight: bold; }
            QLineEdit, QComboBox, QListWidget { 
                background-color: #1C1C26; border: 1px solid #32323D; border-radius: 6px; padding: 8px; color: #F0F0F0; font-size: 14px;
            }
            QPushButton { background: #3A3A4A; color: #E0E0E0; border-radius: 6px; font-weight: bold; }
            #GenerateButton { background: #2980B9; color: white; }
            #GenerateButton:hover { background: #3498DB; }
            #ExitButton:hover { background: #C0392B; color: white; }
            #CapsuleFrame { background-color: rgba(255, 255, 255, 0.03); border: 1px solid #3A3A4A; border-radius: 15px; }
            #GhostBtn_Purple, #GhostBtn_Green { background: transparent; border: none; color: #5A5A6A; font-size: 11px; }
            #GhostBtn_Purple:hover { color: #A29BFE; background-color: rgba(162, 155, 254, 0.1); border-radius: 13px; }
            #GhostBtn_Green:hover { color: #55EFC4; background-color: rgba(85, 239, 196, 0.1); border-radius: 13px; }
        """)


# --- 4. 程序入口 ---
if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle("Fusion")

    if not os.path.exists(CONFIG_FILE):
        QMessageBox.critical(None, "Security Alert", "CRITICAL ERROR: Security config missing. Access Denied.")
        sys.exit()

    s, h = load_config()

    if s == "locked":
        login = LoginWindow(h)
        login.show()
        # 这种方式可以让主界面等待登录窗关闭
        app.exec()

        if login.authorized:
            window = MedicalApp()
            window.show()
            sys.exit(app.exec())
    else:
        window = MedicalApp()
        window.show()
        sys.exit(app.exec())